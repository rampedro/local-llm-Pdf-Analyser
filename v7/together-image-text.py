import os
import logging
import sys
import numpy as np
import json
import chromadb
import ollama
from llama_index.llms.ollama import Ollama
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.core import StorageContext, VectorStoreIndex, Settings, SimpleDirectoryReader, PromptTemplate
from llama_index.vector_stores.chroma import ChromaVectorStore
import fitz  # PyMuPDF for PDFs
import docx  # python-docx for DOCX
from prompts import CATEGORY_EXTRACTION_PROMPT
from sentence_transformers import SentenceTransformer
import torch
from PIL import Image
from io import BytesIO
import base64
import requests

logging.basicConfig(stream=sys.stdout, level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')


# === Document Text Extraction ===
def extract_section(text, section_name):
    start = text.lower().find(section_name)
    if start == -1:
        return ""
    end = text.find("\n", start)
    return text[start:end].strip() if end != -1 else text[start:].strip()


def load_existing_data(doc_dir):
    X, y, filenames = [], [], []
    for fn in os.listdir(doc_dir):
        path = os.path.join(doc_dir, fn)
        text = ""
        if fn.lower().endswith(".txt"):
            with open(path, "r") as f: text = f.read()
        elif fn.lower().endswith(".pdf"):
            try:
                pdf = fitz.open(path)
                for page in pdf:
                    text += page.get_text()
                pdf.close()
            except Exception as e:
                logging.warning(f"Error reading PDF {fn}: {e}")
                continue
        elif fn.lower().endswith(".docx"):
            try:
                doc = docx.Document(path)
                for p in doc.paragraphs:
                    text += p.text
            except Exception as e:
                logging.warning(f"Error reading DOCX {fn}: {e}")
                continue
        else:
            continue

        if text.strip():
            X.append(len(text))
            y.append(len(fn))  # Can be replaced with metadata field
            filenames.append(fn)
    return np.array(X), np.array(y), filenames


# === Image Extraction from PDF ===
def extract_images_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    images = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        img_list = page.get_images(full=True)
        for img_index, img in enumerate(img_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image = Image.open(BytesIO(image_bytes))
            images.append(image)
    return images


def image_to_base64(image):
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    img_str = base64.b64encode(buffer.getvalue()).decode("utf-8")
    return img_str


def send_image_to_model(image_base64):
    url = "http://localhost:11434/api/generate"  # Example URL; replace with correct model endpoint
    payload = {
        "model": "llava",  # Specify your model here
        "prompt": "What is in this picture? explain the visualization in details",
        "images": [image_base64]
    }

    response = requests.post(url, json=payload)

    if response.status_code == 200:
        lines = response.text.splitlines()
        for line in lines:
            try:
                response_json = json.loads(line)
                print("Response description:", response_json["response"])
            except json.JSONDecodeError as e:
                print(f"Failed to decode JSON: {e}")
    else:
        print(f"Error: {response.status_code}")


# === Embedding & Document Processing ===
def get_sentence_level_embedding(text):
    model = SentenceTransformer('paraphrase-MiniLM-L6-v2')
    embedding = model.encode(text)
    return embedding


def init_agent(doc_dir, model_name="BAAI/bge-small-en-v1.5"):
    llm = Ollama(model="granite3.1-dense:2b", request_timeout=300.0)
    embed = HuggingFaceEmbedding(model_name=model_name)
    Settings.llm = llm
    Settings.embed_model = embed

    reader = SimpleDirectoryReader(input_dir=doc_dir, recursive=True)
    docs = reader.load_data()
    logging.info(f"Loaded {len(docs)} documents into memory")

    # Set up ChromaDB for storing document embeddings
    client = chromadb.EphemeralClient()
    collection = client.create_collection("pdf_extract")

    for i, doc in enumerate(docs):
        sentence_embedding = get_sentence_level_embedding(doc.text)
        collection.add(
            ids=[str(i)],
            embeddings=sentence_embedding,
            documents=[doc.text]
        )

    def retrieve_relevant_document(prompt):
        query_embedding = get_sentence_level_embedding(prompt)
        results = collection.query(query_embeddings=[query_embedding], n_results=1)
        data = results['documents'][0][0]
        return data

    def generate_response(prompt, data):
        output = ollama.generate(
            model="granite3.1-dense:2b",
            prompt=f"Using this data: {data}. Respond to this prompt: {prompt}"
        )
        return output['response']

    example_prompt = "Extract high-level information from computer science articles and categorize them."
    retrieved_data = retrieve_relevant_document(example_prompt)
    generated_output = generate_response(example_prompt, retrieved_data)
    logging.info(f"Generated Output: {generated_output}")

    vs = ChromaVectorStore(chroma_collection=collection)
    idx = VectorStoreIndex.from_documents(
        docs,
        storage_context=StorageContext.from_defaults(vector_store=vs),
        embed_model=embed
    )

    tpl = PromptTemplate(CATEGORY_EXTRACTION_PROMPT)
    return idx.as_query_engine(text_qa_template=tpl, similarity_top_k=3)


# === Category Extraction from File ===
def extract_categories_from_file(file_path):
    print(f"Extracting categories from file: {file_path}")
    text = ""

    if file_path.lower().endswith(".pdf"):
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
            doc.close()
        except Exception as e:
            print(f"Error while extracting text from PDF: {e}")
            return
    elif file_path.lower().endswith(".docx"):
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            print(f"Error while extracting text from DOCX: {e}")
            return
    else:
        print("Unsupported file type.")
        return

    if not text.strip():
        print("No text extracted from file.")
        return

    prompt = CATEGORY_EXTRACTION_PROMPT.format(context_str=text[:3000])

    try:
        response = ollama.generate(
            model="granite3.1-dense:2b",
            prompt=prompt
        )
        extracted_info = response["response"]
        print("\n✅ Extracted Categories:\n")
        print(extracted_info)
    except Exception as e:
        print(f"LLM category extraction failed: {e}")


def main():
    while True:
        print("=== PDF/DOCX Category Extraction & Research Helper ===")
        print("1) Extract categories from a file")
        print("0) Exit")
        choice = input("Choose: ")

        if choice == "1":
            file_path = input("Enter the path of the file to extract categories from: ")
            print("Extracting categories from the file...")
            extract_categories_from_file(file_path)
        
        elif choice == "0":
            print("Exiting...")
            break
        
        else:
            print("Invalid option. Please try again.")


if __name__ == "__main__":
    main()
