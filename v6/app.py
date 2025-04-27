import os
import logging
import sys
import numpy as np
import json
import chromadb
import ollama  # Import Ollama for embeddings
from llama_index.llms.ollama import Ollama  # Use the old style LLM import
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.core import StorageContext, VectorStoreIndex, Settings, SimpleDirectoryReader, PromptTemplate
from llama_index.vector_stores.chroma import ChromaVectorStore
import fitz  # PyMuPDF for PDFs
import docx  # python-docx for DOCX
from prompts import CATEGORY_EXTRACTION_PROMPT
from sentence_transformers import SentenceTransformer
import torch

logging.basicConfig(stream=sys.stdout, level=logging.INFO, 
                    format='%(asctime)s - %(levelname)s - %(message)s')

# Function to extract the specific category sections from the document
def extract_section(text, section_name):
    """
    Simple helper to extract text of a specific section from the document.
    E.g., abstract, introduction, conclusion.
    """
    start = text.lower().find(section_name)
    if start == -1:
        return ""
    end = text.find("\n", start)
    return text[start:end].strip() if end != -1 else text[start:].strip()

# Initialize different levels of embeddings
def get_word_level_embedding(text):
    """
    Word-level embeddings using Word2Vec or GloVe.
    """
    # Example using pre-trained GloVe embeddings
    # This can be replaced with your preferred word-level embedding model.
    model = torch.hub.load("nlpjs/word2vec", "glove-wiki-gigaword-100")
    words = text.split()
    embeddings = np.mean([model[word] for word in words if word in model], axis=0)
    return embeddings

def get_sentence_level_embedding(text):
    """
    Sentence-level embeddings using BERT, RoBERTa, or SentenceTransformers.
    """
    # Using SentenceTransformer to get sentence embeddings
    model = SentenceTransformer('paraphrase-MiniLM-L6-v2')
    embedding = model.encode(text)
    return embedding

def get_document_level_embedding(text):
    """
    Document-level embeddings using Universal Sentence Encoder (USE) or Doc2Vec.
    """
    model = SentenceTransformer('all-MiniLM-L6-v2')  # USE-like model
    embedding = model.encode(text)
    return embedding

def get_domain_specific_embedding(text, domain="medical"):
    """
    Domain-specific embeddings, e.g., BioBERT for medical text.
    """
    # Example: Using BioBERT for medical text
    model = SentenceTransformer('biobert-base-cased-v1.1')
    embedding = model.encode(text)
    return embedding

def load_existing_data(doc_dir):
    """
    Pull out text lengths and filenames from all docs in the directory.
    This version adds the ability to extract document sections like abstract, conclusion, etc.
    """
    X, y, filenames, sections = [], [], [], []
    for fn in os.listdir(doc_dir):
        path = os.path.join(doc_dir, fn)
        text = ""
        sections_text = {}
        
        if fn.lower().endswith(".txt"):
            with open(path, "r") as f: text = f.read()
        elif fn.lower().endswith(".pdf"):
            try:
                pdf = fitz.open(path)
                for page in pdf:
                    text += page.get_text()
                pdf.close()
                # Example: Extracting specific sections
                sections_text["abstract"] = extract_section(text, "abstract")
                sections_text["introduction"] = extract_section(text, "introduction")
                sections_text["conclusion"] = extract_section(text, "conclusion")
            except Exception as e:
                logging.warning(f"Error reading PDF {fn}: {e}")
                continue
        elif fn.lower().endswith(".docx"):
            try:
                doc = docx.Document(path)
                for p in doc.paragraphs:
                    text += p.text
                # Example: Extracting sections
                sections_text["abstract"] = extract_section(text, "abstract")
                sections_text["introduction"] = extract_section(text, "introduction")
                sections_text["conclusion"] = extract_section(text, "conclusion")
            except Exception as e:
                logging.warning(f"Error reading DOCX {fn}: {e}")
                continue
        else:
            continue

        if text.strip():
            X.append(len(text))
            y.append(len(fn))  # This can be replaced with metadata field
            filenames.append(fn)
            sections.append(sections_text)
    return np.array(X), np.array(y), filenames, sections

def init_agent(doc_dir, model_name="BAAI/bge-small-en-v1.5"):
    """
    Set up Ollama LLM + vector store using llama_index.
    """
    # LLM setup (old style)
    llm = Ollama(model="granite3.1-dense:2b", request_timeout=300.0)
    embed = HuggingFaceEmbedding(model_name=model_name)
    Settings.llm = llm
    Settings.embed_model = embed

    # Load documents
    reader = SimpleDirectoryReader(input_dir=doc_dir, recursive=True)
    docs   = reader.load_data()
    logging.info(f"Loaded {len(docs)} documents into memory")

    # Set up ChromaDB for storing document embeddings
    client = chromadb.EphemeralClient()
    collection = client.create_collection("pdf_extract")

    # Step 1: Embedding Stage - Embed each document at multiple levels and add to Chroma collection
    for i, doc in enumerate(docs):
        # Word-level embedding
        word_embedding = get_word_level_embedding(doc.text)
        
        # Sentence-level embedding
        sentence_embedding = get_sentence_level_embedding(doc.text)
        
        # Document-level embedding
        doc_embedding = get_document_level_embedding(doc.text)
        
        # Domain-specific embedding (example: medical)
        domain_embedding = get_domain_specific_embedding(doc.text, domain="medical")
        
        # Combine all embeddings into one vector (concatenation or averaging)
        combined_embedding = np.concatenate([word_embedding, sentence_embedding, doc_embedding, domain_embedding])

        # Add combined embedding to the Chroma collection
        collection.add(
            ids=[str(i)], 
            embeddings=combined_embedding, 
            documents=[doc.text]
        )
    
    # Step 2: Retrieval Stage - Use query embedding to retrieve relevant documents
    def retrieve_relevant_document(prompt):
        # Generate an embedding for the query input (sentence-level for simplicity)
        query_embedding = get_sentence_level_embedding(prompt)

        # Retrieve the most relevant document
        results = collection.query(query_embeddings=[query_embedding], n_results=1)
        data = results['documents'][0][0]  # Assuming you want the first document returned
        return data
    
    # Step 3: Generation Stage - Generate response using the LLM with retrieved data
    def generate_response(prompt, data):
        output = ollama.generate(
            model="granite3.1-dense:2b",
            prompt=f"Using this data: {data}. Respond to this prompt: {prompt}"
        )
        return output['response']

    # Example retrieval and generation using CATEGORY_EXTRACTION_PROMPT
    def extract_categories_from_data(data):
        prompt = CATEGORY_EXTRACTION_PROMPT.format(context_str=data)
        response = generate_response(prompt, data)
        return response

    # Example retrieval and generation
    example_prompt = "You are an intelligent assistant. Your task is to extract high-level information from Computer Science articles and categorize it into predefined topics."
    retrieved_data = retrieve_relevant_document(example_prompt)
    generated_output = generate_response(example_prompt, retrieved_data)
    logging.info(f"Generated Output: {generated_output}")

    # Use the CATEGORY_EXTRACTION_PROMPT to extract categories
    categories = extract_categories_from_data(retrieved_data)
    logging.info(f"Extracted Categories: {categories}")

    # Set up the vector store and index documents for general query capabilities
    vs = ChromaVectorStore(chroma_collection=collection)
    idx = VectorStoreIndex.from_documents(
        docs,
        storage_context=StorageContext.from_defaults(vector_store=vs),
        embed_model=embed
    )

    tpl = PromptTemplate(CATEGORY_EXTRACTION_PROMPT)
    return idx.as_query_engine(text_qa_template=tpl, similarity_top_k=3)

# File extraction & category logic
def extract_categories_from_file(file_path):
    print(f"Extracting categories from file: {file_path}")
    text = ""

    # Step 1: Read text from file
    if file_path.lower().endswith(".pdf"):
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
            doc.close()
        except Exception as e:
            print(f"Error while extracting text from PDF: {e}")
            return
    elif file_path.lower().endswith(".docx"):
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            print(f"Error while extracting text from DOCX: {e}")
            return
    else:
        print("Unsupported file type.")
        return

    if not text.strip():
        print("No text extracted from file.")
        return

    # Step 2: Prepare the prompt for the LLM
    prompt = CATEGORY_EXTRACTION_PROMPT.format(context_str=text[:3000])  # optional truncation

    # Step 3: Send prompt to the model and get response
    try:
        response = ollama.generate(
            model="granite3.1-dense:2b",
            prompt=prompt
        )
        extracted_info = response["response"]
        print("\n✅ Extracted Categories:\n")
        print(extracted_info)
    except Exception as e:
        print(f"LLM category extraction failed: {e}")


def main():
    while True:
        print("=== PDF/DOCX Category Extraction & Research Helper ===")
        print("1) Extract categories from a file")
        print("0) Exit")
        choice = input("Choose: ")

        if choice == "1":
            file_path = input("Enter the path of the file to extract categories from: ")
            print("Extracting categories from the file...")
            extract_categories_from_file(file_path)
        
        elif choice == "0":
            print("Exiting...")
            break
        
        else:
            print("Invalid option. Please try again.")


if __name__ == "__main__":
    main()
